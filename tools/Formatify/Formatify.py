# GNU GENERAL PUBLIC LICENSE
# Version 3, 29 June 2007

# Copyright (C) 2024 Moisés Ceñera Fernández

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program. If not, see <https://www.gnu.org/licenses/>.

import wx
import os
from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_path
from fpdf import FPDF
import pypandoc
import pdflatex
from docx import Document
import pandas as pd
from openpyxl import load_workbook
from PIL import Image
import zipfile
import shutil

class MainApp(wx.App):
    def OnInit(self):
        self.frame = FileConverterFrame(None, title="File Converter", size=(800, 600))
        self.frame.Show()
        return True

class FileConverterFrame(wx.Frame):
    def __init__(self, *args, **kw):
        super(FileConverterFrame, self).__init__(*args, **kw)

        panel = wx.Panel(self)
        main_sizer = wx.BoxSizer(wx.VERTICAL)

        instructions = wx.StaticText(panel, label="Seleccione un archivo para ver las opciones de conversión disponibles.")
        instruction_font = wx.Font(12, wx.FONTFAMILY_DEFAULT, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_NORMAL)
        instructions.SetFont(instruction_font)
        main_sizer.Add(instructions, 0, wx.ALL | wx.CENTER, 10)

        # Botón para seleccionar archivo
        select_file_btn = wx.Button(panel, label="📄 Seleccionar Archivo", size=(250, 50))
        select_file_font = wx.Font(12, wx.FONTFAMILY_DEFAULT, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_NORMAL)
        select_file_btn.Bind(wx.EVT_BUTTON, self.on_select_file)
        select_file_btn.SetFont(select_file_font)
        main_sizer.Add(select_file_btn, 0, wx.ALL | wx.CENTER, 10)

        # Sizer para los botones dinámicos de conversión
        self.buttons_sizer = wx.BoxSizer(wx.HORIZONTAL)
        main_sizer.Add(self.buttons_sizer, 0, wx.ALL | wx.CENTER, 10)

        panel.SetSizer(main_sizer)
        self.panel = panel
        self.selected_file_path = None

    def on_select_file(self, event):
        with wx.FileDialog(self, "Seleccione un archivo", wildcard="Archivos (*.pdf;*.docx;*.xlsx;*.jpeg;*.jpg;*.png;*.csv;*.txt)|*.pdf;*.docx;*.xlsx;*.jpeg;*.jpg;*.png;*.csv;*.txt", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST) as file_dialog:
            if file_dialog.ShowModal() == wx.ID_CANCEL:
                return
            self.selected_file_path = file_dialog.GetPath()
            self.show_conversion_options(self.selected_file_path)

    def show_conversion_options(self, filepath):
        # Limpiamos los botones anteriores
        self.buttons_sizer.Clear(True)

        # Identificamos el tipo de archivo y generamos los botones de conversión correspondientes
        if filepath.endswith('.docx'):
            self.add_conversion_button("Convertir Word a PDF", self.convert_docx_to_pdf)
            self.add_conversion_button("Convertir Word a Texto", self.convert_docx_to_text)
        elif filepath.endswith('.xlsx'):
            self.add_conversion_button("Convertir Excel a PDF", self.convert_xlsx_to_pdf)
            self.add_conversion_button("Convertir Excel a CSV", self.convert_xlsx_to_csv)
        elif filepath.endswith('.pdf'):
            self.add_conversion_button("Convertir PDF a Texto", self.convert_pdf_to_text)
            self.add_conversion_button("Convertir PDF a Imagen (JPEG)", self.convert_pdf_to_jpeg)
            self.add_conversion_button("Convertir PDF a Imagen (PNG)", self.convert_pdf_to_png)
        elif filepath.endswith(('.jpeg', '.jpg', '.png')):
            self.add_conversion_button("Convertir Imagen a PDF", self.convert_image_to_pdf)
        elif filepath.endswith('.csv'):
            self.add_conversion_button("Convertir CSV a Excel", self.convert_csv_to_xlsx)
        elif filepath.endswith('.txt'):
            self.add_conversion_button("Convertir Texto a PDF", self.convert_text_to_pdf)

        # Actualizamos la interfaz para mostrar los nuevos botones
        self.panel.Layout()

    def add_conversion_button(self, label, handler):
        button = wx.Button(self.panel, label=label)
        button.Bind(wx.EVT_BUTTON, handler)
        self.buttons_sizer.Add(button, 0, wx.ALL, 5)

    def convert_docx_to_pdf(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".pdf"
        
        # Leer contenido de .docx y escribirlo en un PDF
        doc = Document(self.selected_file_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        
        pdf.output(output_path)
        wx.MessageBox(f"Conversión DOCX a PDF completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)

    def convert_docx_to_text(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".txt"
        
        # Leer el archivo DOCX y guardar el texto en un archivo TXT
        doc = Document(self.selected_file_path)
        with open(output_path, "w") as text_file:
            for para in doc.paragraphs:
                text_file.write(para.text + "\n")
                
        
        wx.MessageBox(f"Conversión DOCX a Texto completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)


    def convert_xlsx_to_pdf(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".pdf"
        
        # Leer contenido de .xlsx y escribirlo en un PDF
        wb = load_workbook(self.selected_file_path)
        sheet = wb.active
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        
        for row in sheet.iter_rows(values_only=True):
            row_data = "  ".join([str(cell) for cell in row if cell is not None])
            pdf.cell(0, 10, row_data, ln=True)
        
        pdf.output(output_path)
        wx.MessageBox(f"Conversión XLSX a PDF completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)

    def convert_xlsx_to_csv(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".csv"
        
        # Convertir el contenido de .xlsx a .csv
        df = pd.read_excel(self.selected_file_path)
        df.to_csv(output_path, index=False)
        
        wx.MessageBox(f"Conversión XLSX a CSV completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)

    def convert_pdf_to_text(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".txt"
        
        # Extraer texto de PDF y guardarlo en un archivo de texto
        with open(self.selected_file_path, 'rb') as file:
            pdf = PdfReader(file)
            with open(output_path, 'w') as text_file:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_file.write(text + "\n")
        
        wx.MessageBox(f"Conversión PDF a Texto completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)

    def convert_pdf_to_jpeg(self, event):
        if not self.selected_file_path:
            return
        output_dir = os.path.splitext(self.selected_file_path)[0] + "_images"
        os.makedirs(output_dir, exist_ok=True)

        # Convertir PDF a imágenes JPEG
        pages = convert_from_path(self.selected_file_path, dpi=200)
        for i, page in enumerate(pages):
            page.save(os.path.join(output_dir, f"page_{i + 1}.jpeg"), "JPEG")

        wx.MessageBox(f"Conversión PDF a JPEG completa. Imágenes guardadas en {output_dir}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)
    
    def convert_pdf_to_png(self, event):
        if not self.selected_file_path:
            return
        output_dir = os.path.splitext(self.selected_file_path)[0] + "_images"
        os.makedirs(output_dir, exist_ok=True)

        # Convertir PDF a imágenes PNG
        pages = convert_from_path(self.selected_file_path, dpi=200)
        for i, page in enumerate(pages):
            page.save(os.path.join(output_dir, f"page_{i + 1}.png"), "PNG")

        wx.MessageBox(f"Conversión PDF a PNG completa. Imágenes guardadas en {output_dir}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)


    def convert_image_to_pdf(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".pdf"
        
        image = Image.open(self.selected_file_path)
        pdf_image = image.convert("RGB")
        pdf_image.save(output_path)

        wx.MessageBox(f"Conversión Imagen a PDF completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)

    def convert_csv_to_xlsx(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".xlsx"
        
        # Leer CSV y guardar en un archivo Excel
        df = pd.read_csv(self.selected_file_path)
        df.to_excel(output_path, index=False)

        wx.MessageBox(f"Conversión CSV a Excel completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)

    def convert_text_to_pdf(self, event):
        if not self.selected_file_path:
            return
        output_path = os.path.splitext(self.selected_file_path)[0] + ".pdf"
        
        # Leer contenido de archivo de texto y escribir en PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        with open(self.selected_file_path, 'r') as file:
            for line in file:
                pdf.cell(0, 10, line.strip(), ln=True)
        
        pdf.output(output_path)
        wx.MessageBox(f"Conversión Texto a PDF completa. Archivo guardado en {output_path}.", "Conversión Exitosa", wx.OK | wx.ICON_INFORMATION)




if __name__ == "__main__":
    app = MainApp()
    app.MainLoop()
